import tkinter as tk
from tkinter import messagebox, filedialog
import os
import shutil
from datetime import datetime, date
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Para alinhamento do parágrafo
from babel.dates import format_date
import logging
import sys


logging.basicConfig(filename='app.log', level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Define o caminho base uma única vez
if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(__file__)
################################################# referenciascampos() ####################################################
def referenciascampos():
    data_atual = date.today()
    data_formatada = format_date(data_atual, format='long', locale='pt_BR')

    n001 = n001_entry.get()
    s001 = s001_entry.get()
    s003 = s003_entry.get()
    s005 = s005_entry.get()
    c001 = c001_entry.get()
    h001 = h001_entry.get()
    d003 = d003_entry.get()
    v001 = v001_entry.get()
    t001 = modo_trabalho_modulo_1.get()

    data_formatada_d002 = datetime.now().strftime("%d/%m/%Y")
    data_formatada_d001 = data_formatada

    referencias = {
        "N001": n001, "S001": s001, "S003": s003, "D001": data_formatada_d001,
        "D002": data_formatada_d002,
        "S005": s005,
        "C001": c001, "H001": h001, "D003": d003, "V001": v001, "T001": t001
    }
    return referencias
######################################## replace_text_in_runs(runs, code, value) #########################################
def replace_text_in_runs(runs, code, value):
    for run in runs:
        run.text = run.text.replace(code, value)

################################################### def novosmodulos(): ##################################################
def novosmodulos():
    referencias = referenciascampos()

    for modulo_idx, module_widgets in enumerate(dynamic_modules_widgets, start=2):
        # Modifica a coleta de valores para incluir o modo de trabalho
        modulo_valores = [widget.get() for widget in module_widgets[:-1] if isinstance(widget, tk.Entry)]
        modo_trabalho = module_widgets[-1].get()  # O último item é a variável StringVar do modo de trabalho

        if len(modulo_valores) == 4:  # Confirma que temos 4 valores de entrada
            c, h, d, v = modulo_valores

            referencias[f'C{modulo_idx:03}'] = c
            referencias[f'H{modulo_idx:03}'] = h
            referencias[f'D{modulo_idx + 2:03}'] = d
            referencias[f'V{modulo_idx:03}'] = v
            referencias[f'T{modulo_idx:03}'] = modo_trabalho
        else:
            logging.error("Módulo não tem todos os 4 valores de entrada")
            break

    return referencias

####################### def replace_marker_with_image(documento, marcador, caminho_imagem, width): #######################
def replace_marker_with_image(documento, marcador, caminho_imagem, width):
    for paragrafo in documento.paragraphs:
        if marcador in paragrafo.text:
            # Limpa o texto do parágrafo removendo o marcador
            texto_limpo = paragrafo.text.replace(marcador, "")
            paragrafo.text = texto_limpo
            # Agora adiciona a imagem ao final do parágrafo
            # Note que isso adicionará a imagem após qualquer texto existente no parágrafo
            paragrafo.add_run().add_picture(caminho_imagem, width=width)

############################################# escolher_local_salvamento(): ###############################################
def escolher_local_salvamento():
    # Obtém o caminho para a pasta de downloads do usuário atual
    downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Define as opções para a caixa de diálogo
    opcoes = {
        'defaultextension': '.docx',
        'filetypes': [('documentos do Word', '.docx')],
        'initialdir': downloads_path,  # Usa o caminho para a pasta de downloads
        'title': 'Salvar documento como...'
    }

    # Abre a caixa de diálogo e guarda o caminho escolhido pelo usuário
    caminho_salvamento = filedialog.asksaveasfilename(**opcoes)

    # Verifica se um caminho foi selecionado (ou seja, o usuário não cancelou a operação)
    if caminho_salvamento:
        messagebox.showinfo("Alerta", f"Caminho escolhido: {caminho_salvamento}")
        return caminho_salvamento
    else:
        messagebox.showinfo("Erro", f"Caminho Inválido ou erro de permissão.")
        # Retorna None ou trata de maneira adequada se o usuário cancelar a operação
        return None


######################################### formatar_nomes_modulos(referencias) ############################################
def formatar_nomes_modulos(referencias):
    # Lista para armazenar os nomes dos módulos
    nomes_modulos = []

    # Adiciona o nome do primeiro módulo
    nomes_modulos.append(referencias["C001"])  # Assumindo que C001 é a chave para o nome do primeiro módulo

    # Adiciona os nomes dos módulos dinâmicos
    modulo_idx = 2
    while f"C{modulo_idx:03}" in referencias:
        nomes_modulos.append(referencias[f"C{modulo_idx:03}"])
        modulo_idx += 1

    # Formata a lista de nomes de módulos para a string final
    if len(nomes_modulos) > 1:
        # Junta todos os nomes com vírgula, exceto o último que é precedido por "e"
        nomes_formatados = ", ".join(nomes_modulos[:-1]) + " e " + nomes_modulos[-1]
    else:
        # Apenas um módulo, então retorna ele
        nomes_formatados = nomes_modulos[0]

    return nomes_formatados

################################# substituir_marcador_modulos(documento, referencias) ####################################
def substituir_marcador_modulos(documento, referencias):
    # Gera a string formatada dos nomes dos módulos
    nomes_modulos_formatados = formatar_nomes_modulos(referencias)
    print("Nomes dos módulos formatados:", nomes_modulos_formatados)

    # Substitui o marcador no documento
    for paragrafo in documento.paragraphs:
        if "C001" in paragrafo.text:
            paragrafo.text = paragrafo.text.replace("C001", nomes_modulos_formatados)

    # Se precisar substituir em células de tabelas também, adicione um loop similar para as células

############################### atualizar_tabela1_com_campos_novos(tabela, referencias) ##################################
def atualizar_tabela1_com_campos_novos(tabela, referencias):
    # Definição dos códigos de campos esperados para cada módulo novo
    codigos_por_modulo = [
        ['C002', 'H002', 'D004', 'V002', 'T002'],
        ['C003', 'H003', 'D005', 'V003', 'T003'],
        ['C004', 'H004', 'D006', 'V004', 'T004'],
        ['C005', 'H005', 'D007', 'V005', 'T005'],
    ]

    # Inicializa o índice da linha a ser preenchida na tabela
    linha_idx = 2  # Começando da terceira linha

    for codigos in codigos_por_modulo:
        config_colunas = []
        todos_codigos_presentes = all(codigo in referencias for codigo in codigos)

        if not all(codigo in referencias for codigo in codigos):
            logging.error("Nem todos os códigos de módulo estão presentes em 'referencias'.")
            break  # Interrompe se algum código estiver faltando
        # Se todos os códigos deste módulo estão presentes em referencias, prepara para preenchimento
        else:
            for codigo in codigos:
                if codigo.startswith('C'):  # Para 'Consultor'
                    config_colunas.append({'codigo': codigo, 'prefixo': 'Consultor ', 'posicao': 'depois'})
                elif codigo.startswith('H'):  # Para 'Horas'
                    config_colunas.append({'codigo': codigo, 'prefixo': 'hs', 'posicao': 'antes'})
                elif codigo.startswith('D'):  # Para 'Data'
                    config_colunas.append({'codigo': codigo, 'prefixo': '', 'posicao': 'depois'})
                elif codigo.startswith('V'):  # Para 'Valor'
                    config_colunas.append({'codigo': '', 'prefixo': 'A combinar', 'posicao': 'depois'})

            # Preenche a linha atual da tabela com os dados configurados
            preencher_linha_tabela(tabela, linha_idx, referencias, config_colunas)

            # Incrementa o índice da linha para a próxima iteração
            linha_idx += 1

########################### preencher_linha_tabela(tabela, linha_idx, referencias, config_colunas) #######################
def preencher_linha_tabela(tabela, linha_idx, referencias, config_colunas):
    if len(tabela.rows) <= linha_idx:
        return

    for col_num, config in enumerate(config_colunas):
        celula = tabela.cell(linha_idx, col_num)
        paragrafo = celula.paragraphs[0]

        if 'codigo' in config and config['codigo']:
            valor = referencias.get(config['codigo'], "")
            if config.get('posicao', '') == 'depois':
                texto = f"{config['prefixo']}{valor}"
            else:
                texto = f"{valor}{config['prefixo']}"
        else:
            texto = config['prefixo']

        # Especial para o caso de consultor + modo de trabalho
        if config['codigo'].startswith('C'):
            # Assume que o código do modo de trabalho tem o mesmo índice que o do consultor, mas com 'T' em vez de 'C'
            codigo_modo = 'T' + config['codigo'][1:]
            modo_trabalho = referencias.get(codigo_modo, "")
            texto = f"{texto} - {modo_trabalho}"

        paragrafo.text = texto
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def atualizar_tabela2_com_campos_novos(tabela, referencias):
    # Assumindo que 'codigos_por_modulo' inclui os códigos para consultar nomes e valores hora
    codigos_por_modulo = [
        ['C002', 'V002'],
        ['C003', 'V003'],
        ['C004', 'V004'],
        ['C005', 'V005'],
        # Adicione mais conforme necessário
    ]

    linha_idx = 2  # Comece a adicionar na primeira linha disponível

    for codigos in codigos_por_modulo:
        if all(codigo in referencias for codigo in codigos):
            c_codigo, v_codigo = codigos
            consultor = referencias[c_codigo]
            valor_hora = referencias[v_codigo]

            # Adiciona uma nova linha se necessário
            if linha_idx >= len(tabela.rows):
                tabela.add_row()

            # Preenche a primeira coluna
            tabela.cell(linha_idx, 0).text = f"Valor hora profissional Consultor {consultor}"

            # Preenche a segunda coluna
            tabela.cell(linha_idx, 1).text = f"R$ {valor_hora}"

            linha_idx += 1
        else:
            logging.error("Nem todos os códigos de módulo estão presentes em 'referencias'.")
            pass


def add_dynamic_module():
    global dynamic_module_count
    if dynamic_module_count >= 4:
        messagebox.showwarning("Limite Atingido", "Número máximo de 5 módulos atingido.")
        logging.info("def add_dynamic_module(): Número máximo de 5 módulos atingido.")
        return

    row_base = 10 + (dynamic_module_count * 2)
    dynamic_module_count += 1

    module_widgets = []

    labels_texts = ["Módulo SAP","Qtde de Horas", "Data Início", "Valor/Hora"]
    coluna_label = 0
    coluna_entry = 1
    for i, text in enumerate(labels_texts):
        label = create_label(inner_frame, f"{text}:", row_base, column=coluna_label)
        coluna_label += 2
        entry = create_entry(inner_frame, row_base, column=coluna_entry)
        coluna_entry += 2
        module_widgets.extend([label, entry])
        logging.info(f"Widgets: {module_widgets}")

    # Adiciona o modo de trabalho como botões de rádio
    modo_trabalho_var = tk.StringVar(value="Remoto")
    for i, mode in enumerate(["Remoto", "Híbrido", "Presencial"]):
        tk.Radiobutton(inner_frame, text=mode, variable=modo_trabalho_var, value=mode, bg=background_color).grid(row=row_base, column=8 + i)

    # Armazena os widgets de entrada e a variável do modo de trabalho para este módulo
    module_widgets.append(modo_trabalho_var)  # Isso armazena a variável, não o widget

    dynamic_modules_widgets.append(module_widgets)  # Adiciona os widgets deste módulo à lista global

    # Atualiza a posição do botão '+' para abaixo do último módulo adicionado
    add_module_button.grid(row=row_base, column=11, pady=(10, 10))

########################################################################################################################
####################################-----Salva o Documento-----#########################################################
def save_document():
    # Capturando os valores dos campos
    global run

    # Define o caminho para a pasta de Downloads do usuário
    # Isso funciona para Windows, macOS e Linux
    #downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')

    # Verifica se algum campo está vazio
    if not all([n001_entry.get().strip(), s001_entry.get().strip(),
                s003_entry.get().strip(), s005_entry.get().strip()]):
        messagebox.showerror("Erro", "Todos os campos devem ser preenchidos.")
        return

    # Combina o caminho da pasta de Downloads com o nome do arquivo
    #documento_final_path = os.path.join(downloads_path, nome_arquivo)
    logging.info("Abrindo diálogo para escolher o local de salvamento")
    try:
        documento_final_path = escolher_local_salvamento()
        if documento_final_path is None:
            raise ValueError("Nenhum caminho de salvamento selecionado pelo usuário.")
        logging.info(f"Caminho de salvamento escolhido: {documento_final_path}")
    except Exception as e:
        logging.error(f"Erro ao escolher o local de salvamento: {e}")
        return

    documento_path = os.path.join(base_path, "Documentos", "Comercial.docx")
    if not documento_path:
        messagebox.showinfo("Erro", f"Problema no caminho de origem do documento")
        logging.error(f"Problema no caminho de origem do documento")
        return None

    # Cria uma cópia do documento original
    documento_copia_path = os.path.join(base_path, "Documentos", "Comercial - Copia.docx")
    try:
        shutil.copy(documento_path, documento_copia_path)
        logging.info(f"Cópia realizada")

    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao criar cópia do documento: {e}")
        logging.error(f"Erro ao criar cópia do documento")
        return

    #Abre a cópia do documento para edição
    try:
        documento = Document(documento_copia_path)
        logging.info(f"Documento de cópia aberto para edição com sucesso")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao abrir o documento para edição: {e}")
        logging.error(f"Erro ao abrir o documento para edição")
        return
    try:
        referencias1 = novosmodulos()
        print("Após novos modulos:", referencias1)
        logging.info(f"Novos módulos dinâmicos adicionados: {referencias1}")
    except Exception as e:
        logging.error(f"Erro ao atribuir as referências dos novos módulos: {str(e)}")
        messagebox.showerror("Erro", f"Erro ao atribuir as referências dos novos módulos: {str(e)}")
        return  # Para garantir que a execução não prossiga em caso de erro

    try:
        referencias2 = referenciascampos()
        if(referencias1 == 0):
            referencias = referencias2
            logging.info(f"Referencias iniciais: {referencias}")
        else:
            referencias = referencias1
            logging.info(f"Referencias campos adicionais: {referencias}")
            substituir_marcador_modulos(documento, referencias)
    except Exception as e:
        logging.error(f"Problema na validação das referências")
        return

    for paragrafo in documento.paragraphs:
        for codigo, valor in referencias.items():
            replace_text_in_runs(paragrafo.runs, codigo, valor)

    imagens_para_adicionar = {
        "{Imagem}": (os.path.join(base_path,"Imagens", "image123.png"), Cm(18)),
        "{logo}": (os.path.join(base_path,"Imagens", "logo.png"), Cm(8)),
        "{Certificações}": (os.path.join(base_path,"Imagens", "certificados.png"), Cm(20)),
        "{projetos}": (os.path.join(base_path,"Imagens","projetos.png"), Cm(14)),
        "{clientes}": (os.path.join(base_path,"Imagens","clientes.png"), Cm(15))
        # Adicionar mais imagens conforme necessário
    }

    for marcador, (caminho_imagem, width) in imagens_para_adicionar.items():
        replace_marker_with_image(documento, marcador, caminho_imagem, width)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for codigo, valor in referencias.items():
                        replace_text_in_runs(paragrafo.runs, codigo, valor)

    tabela1 = documento.tables[1]  # Acessa a segunda tabela do documento
    tabela2 = documento.tables[2]  # Acessa a terceira tabela do documento
    try:
        atualizar_tabela1_com_campos_novos(tabela1, referencias)
        atualizar_tabela2_com_campos_novos(tabela2, referencias)
        logging.info(f"Tabela atualizada com campos novos")
    except Exception as e:
        logging.error(f"Erro ao atualizar tabela com os campos novos: {e}")
    # Salva o documento editado na pasta que o usuario escolher

    try:
        documento.save(documento_final_path)
        messagebox.showinfo("Sucesso", f"Documento salvo com sucesso em: {documento_final_path}")
        logging.info(f"Documento salvo com sucesso em: {documento_final_path}")
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao salvar o documento: {e}, Feche o documento Word gerado anteriormente.")
        logging.error(f"Erro ao salvar o documento")
        return

    # Pergunta ao usuário se deseja encerrar o programa
    if messagebox.askyesno("Encerrar", "Deseja encerrar o programa?"):
        app.destroy()  #Encerra o programa

background_color = "#F1F1F1"
text_color = "#000000"
button_color = "#FFA500"

def center_window(width, height):
    screen_width = main.winfo_screenwidth()
    screen_height = main.winfo_screenheight()
    x = (screen_width / 2) - (width / 2)
    y = (screen_height / 2) - (height / 2)
    main.geometry('%dx%d+%d+%d' % (width, height, x, y))

def create_label(master, text, row, column, pady=5, padx=10, sticky="W"):
    label = tk.Label(master, text=text, bg=background_color, fg=text_color)
    label.grid(row=row, column=column, pady=pady, padx=padx, sticky=sticky)

def create_entry(master, row, column, pady=5, padx=10):
    entry = tk.Entry(master)
    entry.grid(row=row, column=column, pady=pady, padx=padx)
    return entry


def upload_logo():
    # Abre o diálogo de seleção de arquivos para o usuário escolher um arquivo de imagem
    filepath = filedialog.askopenfilename(title="Selecione o arquivo do logo", filetypes=(
    ("PNG files", "*.png"), ("JPEG files", "*.jpg"), ("All files", "*.*")))

    if filepath:  # Verifica se um arquivo foi selecionado
        try:
            # Cria o diretório "Imagens" se não existir
            os.makedirs("Imagens", exist_ok=True)
            # Define o caminho de destino dentro da pasta "Imagens"
            destino = (os.path.join(base_path,"Imagens", "logo.png"))
            # Copia o arquivo selecionado para o destino
            shutil.copy(filepath, destino)
            messagebox.showinfo("Sucesso", "Logo carregado com sucesso na pasta 'Imagens'.")
        except Exception as e:
            messagebox.showerror("Erro", f"Não foi possível carregar o logo: {e}")

dynamic_modules_widgets = []
dynamic_module_count = 0

app = tk.Tk()
# Configura a janela para abrir maximizada
app.state('zoomed')
app.title("Gerador de Proposta T&M")

# Configurando o ícone da janela
app.iconbitmap((os.path.join(base_path,"GUI","icone.ico")))  # Certifique-se de ter o arquivo icone.ico no diretório do projeto
app.configure(bg=background_color)

# Criação do main_frame
main_frame = tk.Frame(app, bg="#F1F1F1")

inner_frame = tk.Frame(main_frame, bg=background_color)
inner_frame.pack(pady=20)  # Isso adiciona um pouco de espaço vertical e centraliza horizontalmente

# Carregando e exibindo o logotipo
logo = tk.PhotoImage(file=(os.path.join(base_path,"GUI","logotipo fusion.png"))) # Assegure-se de que o arquivo esteja no diretório correto
logo_label = tk.Label(app, image=logo, bg=background_color)
logo_label.pack(pady=(50, 20), anchor="n")
main_frame.pack(expand=True, fill='both')

# Crie um frame separado para os botões na parte inferior
buttons_frame = tk.Frame(main_frame, bg=background_color)  # Use main_frame como parent para manter na parte inferior
buttons_frame.pack(side='bottom', fill='x', pady=(0, 150))  # Isso posiciona o frame na parte inferior com padding de 150px no bottom

upload_logo_button = tk.Button(buttons_frame, text="Upload Logo", command=upload_logo, bg=button_color, fg=text_color, cursor="hand2")
upload_logo_button.pack(side='top', pady=5)  # Ajuste o lado conforme necessário

save_button = tk.Button(buttons_frame, text="Salvar Documento", command=save_document, bg=button_color, fg=text_color, cursor="hand2")
save_button.pack(side='top', pady=5)  # Ajuste o lado conforme necessário


#screen.center_window(650, 800)

main_frame = tk.Frame(app, bg=background_color)

create_label(inner_frame, "Número da Proposta:", 0, 0)
create_label(inner_frame, "Nome do Cliente:", 0, 2)
create_label(inner_frame, "Necessidade:", 0, 4)
create_label(inner_frame, "Executivo Comercial:", 0, 6)
create_label(inner_frame, "Módulo SAP:", 1, 0)
create_label(inner_frame, "Qtde de Horas:", 1, 2)
create_label(inner_frame, "Data Início:", 1, 4)
create_label(inner_frame, "Valor/Hora:", 1, 6)

n001_entry = create_entry(inner_frame, 0, 1)
s001_entry = create_entry(inner_frame, 0, 3)
s003_entry = create_entry(inner_frame, 0, 5)
s005_entry = create_entry(inner_frame, 0, 7)
c001_entry = create_entry(inner_frame, 1, 1)
h001_entry = create_entry(inner_frame, 1, 3)
d003_entry = create_entry(inner_frame, 1, 5)
v001_entry = create_entry(inner_frame, 1, 7)

# Adicione o botão para adicionar novos módulos na interface, ajustando sua posição inicial
add_module_button = tk.Button(inner_frame, text="+", command=add_dynamic_module, bg=button_color, fg=text_color, cursor="hand2")
add_module_button.grid(row=1, column=11, pady=(15, 15))

# Valor que representa a escolha do usuário para o modo de trabalho do módulo 1
modo_trabalho_modulo_1 = tk.StringVar(value="Remoto")

# Cria botões de rádio para o modo de trabalho do módulo 1
tk.Radiobutton(inner_frame, text="Remoto", variable=modo_trabalho_modulo_1, value="Remoto", bg=background_color).grid(row=1, column=8)
tk.Radiobutton(inner_frame, text="Híbrido", variable=modo_trabalho_modulo_1, value="Híbrido", bg=background_color).grid(row=1, column=9)
tk.Radiobutton(inner_frame, text="Presencial", variable=modo_trabalho_modulo_1, value="Presencial", bg=background_color).grid(row=1, column=10)




app.mainloop()
